#!/usr/bin/env python3
"""
docx_read.py - 读取 Word 文档（.docx）的工具
功能：
1. 获取文档基本信息：字符数、段落数、表格数
2. 读取指定范围内的段落（支持起始索引和结束索引）
3. 读取指定范围内的表格（支持起始索引和结束索引）
"""

import os
from typing import Dict, Any
from docx import Document

tool_def = {
    "type": "function",
    "function": {
        "name": "docx_read",
        "description": "获取 docx 文件的信息或文字内容。支持：info（获取文档基本信息:字符数、段落数、表格数）、read_paragraphs（读取段落）、read_tables（读取表格）、search（搜索关键字，返回含有关键字的段落编号）。",
        "parameters": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Word 文档的路径（绝对路径或相对于当前工作目录的路径）。"
                },
                "action": {
                    "type": "string",
                    "enum": ["info", "read_paragraphs", "read_tables"],
                    "description": "要执行的操作类型。"
                },
                "start": {
                    "type": "integer",
                    "description": "起始索引（从 0 开始），默认为 0。"
                },
                "end": {
                    "type": "integer",
                    "description": "结束索引（不包含），默认读取从 start 到末尾的内容。"
                },
                "keyword": {
                    "type": "string",
                    "description": "要搜索的关键字，用于 search。"
                }
            },
            "required": ["file_path", "action"]
        }
    }
}

def execute(file_path: str, action: str, start: int = 0, end: int = None, keyword: str = None) -> str:
    """
    执行 docx 读取操作。
    """
    # 检查文件是否存在
    if not os.path.exists(file_path):
        return f"错误：文件 '{file_path}' 不存在。"

    # 尝试加载文档
    try:
        doc = Document(file_path)
    except Exception as e:
        return f"错误：无法打开文档 '{file_path}' - {e}"

    # 根据操作类型分发
    if action == "info":
        return _get_info(doc, file_path)
    elif action == "read_paragraphs":
        return _read_paragraphs(doc, start, end)
    elif action == "read_tables":
        return _read_tables(doc, start, end)
    elif action == "search":
        if not keyword:
            return "错误：搜索操作需要提供 keyword。"
        return _search_paragraphs(doc, keyword, start, end)
    else:
        return f"错误：未知的操作类型 '{action}'。可选值：info, read_paragraphs, read_tables, search。"

def _get_info(doc, file_path: str) -> str:
    """
    获取文档基本信息：字符数、段落数、表格数。
    """
    # 统计字符数：遍历所有段落，计算所有段落文本的字符数
    char_count = 0
    for para in doc.paragraphs:
        char_count += len(para.text)

    # 段落数
    para_count = len(doc.paragraphs)

    # 表格数
    table_count = len(doc.tables)

    info = (
        f"文件：{file_path}\n"
        f"字符数（仅段落文本）：{char_count}\n"
        f"段落数：{para_count}\n"
        f"表格数：{table_count}"
    )
    return info

def _read_paragraphs(doc, start: int, end: int = None) -> str:
    """
    读取指定范围内的段落。
    """
    total = len(doc.paragraphs)
    if total == 0:
        return "文档中没有段落。"

    # 处理索引边界
    if start < 0:
        start = 0
    if start >= total:
        return f"起始索引 {start} 超出范围（最大索引 {total-1}）。"

    if end is None:
        end = total
    else:
        if end < 0:
            end = 0
        if end > total:
            end = total
    if start >= end:
        start,end=end,start

    # 提取段落内容
    paragraphs = []
    for i in range(start, end):
        para_text = doc.paragraphs[i].text.strip()
        if not para_text:
            para_text = "(空段落)"
        paragraphs.append(f"[{i}] {para_text}")

    return "\n".join(paragraphs)

def _read_tables(doc, start: int, end: int = None) -> str:
    """
    读取指定范围内的表格。
    """
    total = len(doc.tables)
    if total == 0:
        return "文档中没有表格。"

    # 处理索引边界
    if start < 0:
        start = 0
    if start >= total:
        return f"起始索引 {start} 超出范围（最大索引 {total-1}）。"

    if end is None:
        end = total
    else:
        if end < 0:
            end = 0
        if end > total:
            end = total
    if start >= end:
        start,end=end,start

    # 提取表格内容，每个表格格式化为 Markdown 表格（或简单文本）
    output_lines = []
    for i in range(start, end):
        table = doc.tables[i]
        output_lines.append(f"表格 {i}:")
        # 将表格转换为文本表示
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            output_lines.append("\n".join(rows))
        else:
            output_lines.append("(空表格)")
        output_lines.append("")  # 空行分隔
    return "\n".join(output_lines)
def _search_paragraphs(doc, keyword: str, start: int = 0, end: int = None) -> str:
    """
    在指定范围内搜索包含关键字的段落，返回段落编号列表。
    """
    total = len(doc.paragraphs)
    if total == 0:
        return "文档中没有段落。"

    # 处理索引边界
    if start < 0:
        start = 0
    if start >= total:
        return f"起始索引 {start} 超出范围（最大索引 {total-1}）。"

    if end is None:
        end = total
    else:
        if end < 0:
            end = 0
        if end > total:
            end = total
    if start >= end:
        start,end=end,start

    matched_paragraphs = []
    for i in range(start, end):
        if keyword in doc.paragraphs[i].text:
            matched_paragraphs.append(i)

    if not matched_paragraphs:
        return f"未找到包含关键字 '{keyword}' 的段落。"
    return f"包含关键字 '{keyword}' 的段落编号：{matched_paragraphs}"